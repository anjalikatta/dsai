# generate_homework1.py
# Generates the HOMEWORK1 .docx submission for the AI-Powered Reporter Software
# Run: python 03_query_ai/generate_homework1.py

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)

GITHUB_BASE = "https://github.com/anjalikatta/dsai/blob/main"

# ─── Helper functions ───

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

def add_body(text):
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    return p

def add_bold_body(bold_text, normal_text=""):
    p = doc.add_paragraph()
    run = p.add_run(bold_text)
    run.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_link_line(label, url):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(f"{label}: ")
    run.bold = True
    p.add_run(url)

def add_screenshot_placeholder(caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[INSERT SCREENSHOT: {caption}]")
    run.italic = True
    run.font.color.rgb = RGBColor(180, 0, 0)

# ═══════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════

title = doc.add_heading("Homework 1: AI-Powered Reporter Software", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("SYSEN 5381 — Data Science and AI for Systems Engineering")
run.font.size = Pt(14)

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Anjali Katta").bold = True
info.add_run("\nFDA Medical Device Recall Reporter")

doc.add_paragraph()
doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION 1: WRITING COMPONENT (25 pts)
# ═══════════════════════════════════════════════════════════

add_heading("1. System Overview", level=1)

add_body(
    "My AI-powered reporter software integrates three components I built over the past three weeks "
    "into a single pipeline that retrieves FDA medical device recall data, displays it through an "
    "interactive web interface, and generates analytical reports using artificial intelligence. "
    "The project centers on the FDA's open Device Recall API, which provides information about "
    "medical devices that have been recalled due to safety or quality concerns. I chose this API "
    "because medical device safety is directly relevant to systems engineering — understanding "
    "failure modes, root causes, and trends helps engineers design better products and quality "
    "systems."
)

add_body(
    "The system works as a three-stage pipeline. First, an R script (my_good_query.R) handles the "
    "raw API query — it loads an API key from a .env file, sends a GET request to the FDA endpoint "
    "with date-range and limit parameters, parses the JSON response, and prints a sample of the "
    "results. Second, a Shiny web application (app.R and utils.R) wraps that same API logic in a "
    "browser-based interface where users can pick start and end dates, choose how many records to "
    "retrieve, and explore the results in a sortable, filterable data table. The app uses the bslib "
    "package for a modern UI theme and includes status messages and error handling so users know "
    "exactly what's happening. Third, a Python script (ai_reporter.py) queries the same FDA "
    "endpoint, aggregates the data into summary statistics — top root causes, monthly recall "
    "counts — and sends that summary to OpenAI's GPT-4o-mini model with a carefully designed "
    "prompt to produce a structured Markdown report."
)

add_body(
    "The biggest design choice was how to format data for the AI prompt. In early iterations, "
    "I sent raw JSON to the model and the output was vague and generic. By pre-aggregating the "
    "data into root-cause counts and monthly totals, then embedding those summaries directly in "
    "the prompt, the AI could reference specific numbers and produce a much more useful report. "
    "I also iterated on the prompt itself three times — starting broad, then adding explicit section "
    "headers, and finally specifying exact sentence and bullet counts — to get consistent, "
    "professional output. One challenge was handling the FDA's date format (YYYYMMDD strings) "
    "across both R and Python; I wrote formatting helpers in both languages to convert them to "
    "readable YYYY-MM-DD dates. Overall, combining these three components showed me how API "
    "queries, interactive applications, and AI generation can work together in a real data "
    "pipeline."
)

p_note = doc.add_paragraph()
run = p_note.add_run(
    "NOTE TO STUDENT: The writing component above is a draft. The homework requires this section "
    "to be written in your own words (NOT AI-generated). Please review, revise, and personalize "
    "this text before submitting."
)
run.italic = True
run.font.color.rgb = RGBColor(180, 0, 0)
run.font.size = Pt(9)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION 2: GIT REPOSITORY LINKS (25 pts)
# ═══════════════════════════════════════════════════════════

add_heading("2. Git Repository Links", level=1)

add_body("All source code is hosted on GitHub. Links below point to the specific files on the main branch.")

add_heading("API Query Script", level=2)
add_link_line(
    "my_good_query.R",
    f"{GITHUB_BASE}/01_query_api/my_good_query.R"
)
add_link_line(
    "README (query documentation)",
    f"{GITHUB_BASE}/01_query_api/README_my_good_query.md"
)

add_heading("Shiny Application", level=2)
add_link_line(
    "app.R (main Shiny app)",
    f"{GITHUB_BASE}/02_productivity/shiny_app/app.R"
)
add_link_line(
    "utils.R (helper functions)",
    f"{GITHUB_BASE}/02_productivity/shiny_app/utils.R"
)
add_link_line(
    "DESCRIPTION (R dependencies)",
    f"{GITHUB_BASE}/02_productivity/shiny_app/DESCRIPTION"
)
add_link_line(
    "README (Shiny app documentation)",
    f"{GITHUB_BASE}/02_productivity/shiny_app/README.md"
)

add_heading("AI Reporter Script", level=2)
add_link_line(
    "ai_reporter.py (AI-powered reporter)",
    f"{GITHUB_BASE}/03_query_ai/ai_reporter.py"
)
add_link_line(
    "lab_ai_reporter_output.md (generated report)",
    f"{GITHUB_BASE}/03_query_ai/lab_ai_reporter_output.md"
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION 3: SCREENSHOTS / OUTPUTS (25 pts)
# ═══════════════════════════════════════════════════════════

add_heading("3. Screenshots and Outputs", level=1)

add_heading("3.1 Shiny App — Interface", level=2)
add_body("The Shiny app displays query parameter controls (date range, max records) and a Query button:")
add_screenshot_placeholder("Shiny app interface with date inputs and Query FDA API button")

add_heading("3.2 Shiny App — Query Results", level=2)
add_body("After clicking Query FDA API, the app displays results in a sortable, filterable DataTable:")
add_screenshot_placeholder("Shiny app showing recall data in the results table")

add_heading("3.3 Shiny App — Error Handling", level=2)
add_body("When the start date is after the end date, the app shows an error message:")
add_screenshot_placeholder("Shiny app showing error handling (e.g., invalid date range)")

add_heading("3.4 AI-Generated Report (Sample Output)", level=2)
add_body(
    "Below is the full AI-generated report produced by ai_reporter.py using OpenAI GPT-4o-mini. "
    "This report was saved to lab_ai_reporter_output.md."
)

report_text = """# FDA Medical Device Recall Report (2024)

## Executive Summary
In 2024, the FDA recorded a total of 1,000 medical device recalls, reflecting significant concerns in device safety and quality assurance. Notably, September emerged as the peak month for recalls, suggesting vulnerabilities in manufacturing and reporting processes during this period.

## Key Findings
- The leading cause of recalls was "Under Investigation by firm," contributing to 27% (270 recalls) of all incidents.
- Process control failures accounted for 19% (188 recalls), emphasizing the need for stringent quality checks in production.
- A notable spike in recalls occurred in September (271), which saw over 27% of the total recalls for the year.
- Monthly recall counts were exceptionally low in the earlier months (January to February) but escalated dramatically after mid-year, indicating potential lapses in manufacturing practices or oversight.

## Root Cause Analysis
The dominant causes of recalls were largely centered around investigative issues and process control, with "Under Investigation by firm" and "Process control" together constituting nearly half of all recalls (458). This suggests that companies may be struggling with internal quality management or are potentially reactive rather than proactive in addressing device safety concerns.

## Monthly Trends
Recall volumes exhibited significant variability throughout the year, with a marked increase beginning mid-year. After January's low count of 20 recalls, a dramatic rise was observed in September (271 recalls), indicating possible systemic issues or failures in manufacturing processes that required urgent attention.

## Recommendations
- Enhance Quality Assurance Protocols: Manufacturers should invest in robust process controls and regular audits to proactively identify potential risks before they result in recalls.
- Increase Transparency: Firms should foster greater transparency with the FDA regarding their quality assurance practices and recall responses to facilitate more effective oversight.
- Focus on Design and Component Selection: Particular attention should be given to the design phases and selection of components to mitigate the risk of recalls due to design flaws or nonconformance."""

for line in report_text.strip().split("\n"):
    line = line.strip()
    if not line:
        doc.add_paragraph()
    elif line.startswith("# "):
        h = doc.add_heading(line[2:], level=3)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 80, 120)
    elif line.startswith("## "):
        h = doc.add_heading(line[3:], level=4)
        for run in h.runs:
            run.font.color.rgb = RGBColor(0, 80, 120)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    else:
        doc.add_paragraph(line)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# SECTION 4: DOCUMENTATION (25 pts)
# ═══════════════════════════════════════════════════════════

add_heading("4. Documentation", level=1)

# --- 4.1 Data Summary ---
add_heading("4.1 Data Summary", level=2)
add_body(
    "The FDA Device Recall API returns JSON objects with many fields. "
    "The table below summarizes the key columns used in this project."
)

data_cols = [
    ("recall_number",           "character (string)", "Unique FDA recall identifier (e.g., Z-1234-2024)"),
    ("event_date_initiated",    "character / date",   "Date the recall was initiated; raw format YYYYMMDD, converted to YYYY-MM-DD in the app"),
    ("product_code",            "character (string)", "FDA product classification code for the recalled device"),
    ("root_cause_description",  "character (string)", "Text description of the root cause of the recall (e.g., Process Control, Software Design)"),
]

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = table.rows[0].cells
hdr[0].text = "Column Name"
hdr[1].text = "Data Type"
hdr[2].text = "Description"
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True

for name, dtype, desc in data_cols:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = dtype
    row[2].text = desc

doc.add_paragraph()

add_body(
    "Additional fields available from the API (not displayed by default): "
    "recalling_firm, product_description, code_info, cfres_id, res_event_number, "
    "k_numbers, openfda.device_name, openfda.medical_specialty_description."
)

# --- 4.2 Technical Details ---
add_heading("4.2 Technical Details", level=2)

add_bold_body("API Endpoint: ", "https://api.fda.gov/device/recall.json (GET)")
add_bold_body("API Authentication: ", "API key passed as query parameter. Free key from https://open.fda.gov/apis/authentication/")
add_bold_body("AI Model: ", "OpenAI GPT-4o-mini via https://api.openai.com/v1/chat/completions")

add_heading("Required API Keys", level=3)
env_table = doc.add_table(rows=1, cols=3)
env_table.style = "Light Grid Accent 1"
env_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = env_table.rows[0].cells
hdr[0].text = "Variable"
hdr[1].text = "Location"
hdr[2].text = "Purpose"
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True
for var, loc, purpose in [
    ("API_KEY",        "01_query_api/.env", "FDA Open API key (for higher rate limits)"),
    ("OPENAI_API_KEY", ".env (project root)", "OpenAI API key (required for AI report generation)"),
]:
    row = env_table.add_row().cells
    row[0].text = var
    row[1].text = loc
    row[2].text = purpose

doc.add_paragraph()

add_heading("Project File Structure", level=3)
files_info = [
    ("01_query_api/my_good_query.R",           "R script — queries FDA API, prints 15 sample rows"),
    ("01_query_api/README_my_good_query.md",    "Documentation for the API query script"),
    ("02_productivity/shiny_app/app.R",         "R Shiny app — interactive web UI for FDA recalls"),
    ("02_productivity/shiny_app/utils.R",       "Helper functions (API client, date formatting)"),
    ("02_productivity/shiny_app/DESCRIPTION",   "R package dependencies for the Shiny app"),
    ("02_productivity/shiny_app/README.md",     "Documentation for the Shiny app"),
    ("03_query_ai/ai_reporter.py",             "Python script — queries API, sends to OpenAI, saves report"),
    ("03_query_ai/lab_ai_reporter_output.md",  "Sample AI-generated report output"),
]
file_table = doc.add_table(rows=1, cols=2)
file_table.style = "Light Grid Accent 1"
file_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = file_table.rows[0].cells
hdr[0].text = "File"
hdr[1].text = "Description"
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True
for filepath, desc in files_info:
    row = file_table.add_row().cells
    row[0].text = filepath
    row[1].text = desc

doc.add_paragraph()

add_heading("R Packages (Shiny App)", level=3)
r_pkg_table = doc.add_table(rows=1, cols=2)
r_pkg_table.style = "Light Grid Accent 1"
r_pkg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = r_pkg_table.rows[0].cells
hdr[0].text = "Package"
hdr[1].text = "Purpose"
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True
for pkg, purpose in [
    ("shiny",    "Web application framework"),
    ("bslib",    "Bootstrap theme and layout"),
    ("httr",     "HTTP requests to FDA API"),
    ("jsonlite", "JSON parsing"),
    ("dplyr",    "Data manipulation"),
    ("DT",       "Interactive DataTable widget"),
]:
    row = r_pkg_table.add_row().cells
    row[0].text = pkg
    row[1].text = purpose

doc.add_paragraph()

add_heading("Python Packages (AI Reporter)", level=3)
py_pkg_table = doc.add_table(rows=1, cols=2)
py_pkg_table.style = "Light Grid Accent 1"
py_pkg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = py_pkg_table.rows[0].cells
hdr[0].text = "Package"
hdr[1].text = "Purpose"
for h in hdr:
    for p in h.paragraphs:
        for r in p.runs:
            r.bold = True
for pkg, purpose in [
    ("requests",      "HTTP requests to FDA and OpenAI APIs"),
    ("pandas",        "Data manipulation and aggregation"),
    ("python-dotenv", "Load .env files for API keys"),
]:
    row = py_pkg_table.add_row().cells
    row[0].text = pkg
    row[1].text = purpose

doc.add_paragraph()

# --- 4.3 Usage Instructions ---
add_heading("4.3 Usage Instructions", level=2)

add_heading("Prerequisites", level=3)
for item in [
    "R (>= 4.0) with packages: shiny, bslib, httr, jsonlite, dplyr, DT",
    "Python 3.11+ with packages: requests, pandas, python-dotenv",
    "FDA API key (free from https://open.fda.gov/apis/authentication/)",
    "OpenAI API key (from https://platform.openai.com/api-keys)",
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading("Step 1: Clone the Repository", level=3)
p = doc.add_paragraph()
run = p.add_run("git clone https://github.com/anjalikatta/dsai.git && cd dsai")
run.font.name = "Consolas"
run.font.size = Pt(10)

add_heading("Step 2: Set Up API Keys", level=3)
add_body("Create two .env files:")
for item in [
    '01_query_api/.env — add the line: API_KEY=your_fda_api_key_here',
    '.env (project root) — add the line: OPENAI_API_KEY=your_openai_key_here',
]:
    doc.add_paragraph(item, style="List Bullet")

add_heading("Step 3: Install R Dependencies", level=3)
p = doc.add_paragraph()
run = p.add_run('install.packages(c("shiny", "bslib", "httr", "jsonlite", "dplyr", "DT"))')
run.font.name = "Consolas"
run.font.size = Pt(10)

add_heading("Step 4: Install Python Dependencies", level=3)
p = doc.add_paragraph()
run = p.add_run("pip install requests pandas python-dotenv")
run.font.name = "Consolas"
run.font.size = Pt(10)

add_heading("Step 5: Run the API Query Script", level=3)
p = doc.add_paragraph()
run = p.add_run("Rscript 01_query_api/my_good_query.R")
run.font.name = "Consolas"
run.font.size = Pt(10)
add_body("This prints the first 15 FDA device recall records for 2024 to the console.")

add_heading("Step 6: Run the Shiny App", level=3)
p = doc.add_paragraph()
run = p.add_run('R -e \'shiny::runApp("02_productivity/shiny_app")\'')
run.font.name = "Consolas"
run.font.size = Pt(10)
add_body("Open the URL printed in the console (typically http://127.0.0.1:xxxx). "
         "Set date range, choose max records, and click Query FDA API.")

add_heading("Step 7: Run the AI Reporter", level=3)
p = doc.add_paragraph()
run = p.add_run("python 03_query_ai/ai_reporter.py")
run.font.name = "Consolas"
run.font.size = Pt(10)
add_body("The script queries the FDA API, processes the data, sends it to OpenAI, "
         "and saves the report to 03_query_ai/lab_ai_reporter_output.md.")

# ─── Save ───

output_path = os.path.join(os.path.dirname(__file__), "HOMEWORK1_submission.docx")
doc.save(output_path)
print(f"Saved: {output_path}")
