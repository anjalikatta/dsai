# ai_reporter.py
# AI-Powered FDA Device Recall Reporter
# Pairs with LAB_ai_reporter.md

# This script queries the FDA Device Recall API for 2024 data,
# processes and summarizes the results, then uses OpenAI to
# generate an analytical report saved in multiple formats.

# 0. SETUP ###################################

## 0.1 Load Packages #################################

# pip install requests pandas python-dotenv markdown python-docx

import sys
import requests
import json
import os
import pandas as pd
from dotenv import load_dotenv
import markdown
from docx import Document

# Fix Windows terminal encoding for emoji/unicode output
sys.stdout.reconfigure(encoding="utf-8")

## 0.2 Load Environment Variables #####################

# Load both .env files: root has OPENAI_API_KEY, 01_query_api has FDA API_KEY
load_dotenv(os.path.join(os.path.dirname(__file__), "..", ".env"))
load_dotenv(os.path.join(os.path.dirname(__file__), "..", "01_query_api", ".env"))

FDA_API_KEY = os.getenv("API_KEY")
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

if not FDA_API_KEY:
    raise ValueError("API_KEY not found. Check 01_query_api/.env")
if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY not found. Check root .env")

# 1. QUERY FDA API ###################################

print("📡 Querying FDA API for 2024 Device Recalls...")

base_url = "https://api.fda.gov/device/recall.json"
params = {
    "api_key": FDA_API_KEY,
    "search": "event_date_initiated:[2024-01-01 TO 2024-12-31]",
    "limit": 1000
}

response = requests.get(base_url, params=params)
response.raise_for_status()

data = response.json()
recalls = pd.DataFrame(data["results"])

print(f"✅ Retrieved {len(recalls)} recall records")

# 2. PROCESS DATA ###################################

# Select key columns available in the dataset
desired_cols = [
    "recall_number", "event_date_initiated",
    "product_code", "root_cause_description"
]
cols = [c for c in desired_cols if c in recalls.columns]
df = recalls[cols].copy()

# Parse dates and extract month names for trend analysis
df["event_date_initiated"] = pd.to_datetime(df["event_date_initiated"])
df["month"] = df["event_date_initiated"].dt.month_name()

# Aggregate summary statistics for the AI prompt
total_recalls = len(df)
root_cause_counts = df["root_cause_description"].value_counts().head(10)
monthly_counts = (df
    .groupby(df["event_date_initiated"].dt.month)
    .size()
    .reset_index(name="count")
)
monthly_counts["month_name"] = pd.to_datetime(monthly_counts["event_date_initiated"], format="%m").dt.month_name()

# Format processed data as structured text for the AI
data_summary = f"""FDA Medical Device Recall Data (January - December 2024):
- Total recalls retrieved: {total_recalls}

Top 10 Root Causes of Recall:
{root_cause_counts.to_string()}

Monthly Recall Counts:
{monthly_counts[['month_name', 'count']].to_string(index=False)}

Sample Records (first 5):
{df.head(5).to_string(index=False)}
"""

print("\n📊 Data processed. Sending to OpenAI for analysis...\n")

# 3. GENERATE AI REPORT ###################################

# Prompt design (iteration 3 - final version):
#   v1: "Summarize this FDA recall data." -> Too vague, output was generic
#   v2: Added section structure -> Better, but inconsistent length/format
#   v3: Added specific sentence counts, bullet counts, and tone guidance
#       -> Reliable, structured, professional output

prompt = f"""You are a data analyst preparing a report on FDA medical device recalls in 2024.

Analyze this data and write a structured report:

{data_summary}

Report requirements:
1. **Executive Summary**: 2-3 sentence overview of the recall landscape
2. **Key Findings**: Exactly 4 bullet points highlighting the most important patterns
3. **Root Cause Analysis**: 2-3 sentences on the dominant causes and what they suggest
4. **Monthly Trends**: 2-3 sentences on how recall volume changed throughout the year
5. **Recommendations**: 3 actionable bullet points for device manufacturers or regulators

Format as Markdown with ## headers for each section. Be specific and reference actual numbers from the data. Keep the tone professional and concise."""

url = "https://api.openai.com/v1/chat/completions"

body = {
    "model": "gpt-4o-mini",
    "messages": [
        {"role": "user", "content": prompt}
    ]
}

headers = {
    "Authorization": f"Bearer {OPENAI_API_KEY}",
    "Content-Type": "application/json"
}

ai_response = requests.post(url, headers=headers, json=body)
ai_response.raise_for_status()

result = ai_response.json()
report_text = result["choices"][0]["message"]["content"]

print("📝 AI Report Generated:\n")
print(report_text)

# 4. SAVE REPORT IN MULTIPLE FORMATS ###################################

## 4.1 Save as Markdown (.md) ########################

with open("report.md", "w", encoding="utf-8") as f:
    f.write(report_text)
print("\n✅ Saved report.md")

## 4.2 Save as Plain Text (.txt) #####################

with open("report.txt", "w", encoding="utf-8") as f:
    f.write(report_text)
print("✅ Saved report.txt")

## 4.3 Save as HTML (.html) ##########################

html_content = markdown.markdown(report_text)
html_document = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>FDA Device Recall Report 2024</title>
    <style>
        body {{ font-family: Arial, sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; line-height: 1.6; }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #34495e; margin-top: 30px; }}
        li {{ margin-bottom: 6px; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

with open("report.html", "w", encoding="utf-8") as f:
    f.write(html_document)
print("✅ Saved report.html")

## 4.4 Save as Word Document (.docx) #################

doc = Document()
for line in report_text.split("\n"):
    if line.startswith("# "):
        doc.add_heading(line[2:], level=1)
    elif line.startswith("## "):
        doc.add_heading(line[3:], level=2)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style="List Bullet")
    elif line.strip():
        doc.add_paragraph(line)

doc.save("report.docx")
print("✅ Saved report.docx")

print("\n✅ All report formats saved successfully!")
