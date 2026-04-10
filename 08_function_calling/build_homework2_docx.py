"""Build the Homework 2 .docx submission file for Anjali Katta."""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

REPO = "https://github.com/anjalikatta/dsai/blob/main"

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic

def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def link_bullet(label, path):
    url = f"{REPO}/{path}"
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label)
    run.bold = True
    p.add_run(f": {url}")

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)

def spacer():
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────

title = doc.add_heading("Homework 2: AI Agent System with RAG and Tools", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

para("Anjali Katta", bold=True)
para("FDA Device Recall AI Agent System", italic=True)
spacer()

# ══════════════════════════════════════════════════════════
# SECTION 1: WRITING COMPONENT (25 pts)
# ══════════════════════════════════════════════════════════

heading("1. Writing Component")

doc.add_paragraph(
    "This project builds an AI agent system that analyzes FDA medical device "
    "recalls by combining three techniques learned over the past three weeks: "
    "multi-agent orchestration, retrieval-augmented generation (RAG), and function "
    "calling with tools. The system is built in Python and uses a local "
    "Ollama LLM (smollm2:1.7b) for all agent interactions. My earlier labs each "
    "explored these techniques individually — a multi-agent pipeline for FDA drug "
    "shortage analysis, a RAG system for systems engineering risk advice, and now "
    "a function-calling workflow for FDA device recalls — and this homework "
    "compiles them into one cohesive system."
)

doc.add_paragraph(
    "The unified system works as a three-agent pipeline. Agent 1 (the Data Fetcher) "
    "uses function calling to query the openFDA Device Recall API and retrieve live "
    "recall records for a given year. The LLM decides the tool arguments based on "
    "the task prompt, calls the get_fda_recalls() function, and returns a structured "
    "DataFrame. Agent 2 (the Context Researcher) uses RAG to search a local "
    "knowledge base I created — a text file containing FDA domain expertise about "
    "recall classification levels (Class I/II/III), common root causes, the "
    "regulatory process, and patient safety impact. This gives the system context "
    "that raw API data alone cannot provide. Agent 3 (the Executive Reporter) "
    "receives both the live data from Agent 1 and the synthesized context from "
    "Agent 2, then writes a professional executive brief with sections for overview, "
    "key findings, regulatory context, and recommendations."
)

doc.add_paragraph(
    "The main design challenge was making the three components work together "
    "coherently. For function calling, the tool function had to be injected into "
    "the functions module's namespace so Ollama's tool dispatch could find it. "
    "For RAG, I chose a text-based knowledge base because the domain context is "
    "naturally unstructured (paragraphs of regulatory explanations), and a simple "
    "line-based substring search was sufficient given the focused scope. The "
    "multi-agent chaining uses the agent_run() helper from the course's functions.py, "
    "passing each agent's output as the next agent's task input. A limitation is "
    "that smollm2:1.7b is a small model so its outputs can be imprecise, but the "
    "pipeline architecture is model-agnostic and would produce stronger results "
    "with a larger model."
)

spacer()

# ══════════════════════════════════════════════════════════
# SECTION 2: GIT REPOSITORY LINKS (25 pts)
# ══════════════════════════════════════════════════════════

heading("2. Code — Git Repository Links")

para("All code is in the dsai repository on GitHub:")

link_bullet("Multi-agent orchestration script (Lab 1)",
            "06_agents/lab_prompt_design.py")

link_bullet("RAG implementation (Lab 2)",
            "07_rag/lab_custom_rag_query.py")

link_bullet("RAG knowledge base data (FDA domain context)",
            "07_rag/data/fda_recall_knowledge.txt")

link_bullet("Function calling / tool definitions (Lab 3)",
            "08_function_calling/lab_multi_agent_with_tools.py")

link_bullet("Main system file (Homework 2 — all components combined)",
            "08_function_calling/homework2_fda_agent.py")

link_bullet("Documentation",
            "08_function_calling/DOCUMENTATION.md")

link_bullet("Helper functions (course-provided, used by all scripts)",
            "08_function_calling/functions.py")

spacer()

# ══════════════════════════════════════════════════════════
# SECTION 3: SCREENSHOTS / OUTPUTS (25 pts)
# ══════════════════════════════════════════════════════════

heading("3. Screenshots / Outputs")

para("Note: Insert screenshots from running each script in the terminal.",
     italic=True)
spacer()

heading("3.1 Multi-Agent Workflow (lab_prompt_design.py)", level=2)
para("This screenshot shows three agents chaining together: Agent 1 summarizes "
     "FDA drug shortage data, Agent 2 assesses risks, and Agent 3 writes a "
     "public health advisory.", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 06_agents && python3 lab_prompt_design.py']",
     bold=True)
spacer()

heading("3.2 RAG Retrieval and Response (lab_custom_rag_query.py)", level=2)
para("This screenshot shows the RAG workflow: searching the project risks CSV "
     "and FDA knowledge base, then passing retrieved context to the LLM for "
     "context-aware answers.", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 07_rag && python3 lab_custom_rag_query.py']",
     bold=True)
spacer()

heading("3.3 Function Calling / Tool Usage (lab_multi_agent_with_tools.py)", level=2)
para("This screenshot shows Agent 1 calling the get_fda_recalls() tool to "
     "fetch live data from the openFDA API, returning a DataFrame of recall "
     "records. Agent 2 then analyzes the data.", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 08_function_calling && python3 lab_multi_agent_with_tools.py']",
     bold=True)
spacer()

heading("3.4 Combined System (homework2_fda_agent.py)", level=2)
para("This screenshot shows the unified system with all three components: "
     "function calling (Agent 1), RAG (Agent 2), and the final executive "
     "report (Agent 3).", italic=True)
para("[INSERT SCREENSHOT: Run 'cd 08_function_calling && python3 homework2_fda_agent.py']",
     bold=True)
spacer()

# ══════════════════════════════════════════════════════════
# SECTION 4: DOCUMENTATION (25 pts)
# ══════════════════════════════════════════════════════════

heading("4. Documentation")

heading("4.1 System Architecture", level=2)

para("The system uses a 3-agent pipeline where each agent has a specific role:")

add_table(
    ["Agent", "Role", "Component", "Input", "Output"],
    [
        ["Agent 1: Data Fetcher",
         "Queries openFDA API",
         "Function Calling",
         "LLM picks tool args (year, limit)",
         "DataFrame of recall records"],
        ["Agent 2: Context Researcher",
         "Searches local knowledge base",
         "RAG",
         "Search results for root causes, classifications, safety",
         "Synthesized domain context"],
        ["Agent 3: Executive Reporter",
         "Writes professional report",
         "Multi-Agent",
         "Data from Agent 1 + context from Agent 2",
         "Executive brief (Overview, Findings, Context, Recommendations)"],
    ],
)
spacer()

heading("4.2 RAG Data Source", level=2)
bullet("File: 07_rag/data/fda_recall_knowledge.txt")
bullet("Type: Plain text (~40 lines) of FDA domain knowledge")
bullet("Search: Case-insensitive substring matching across all lines")
bullet("Topics: Recall classifications, root causes, regulatory process, patient safety")
spacer()

heading("4.3 Tool Functions", level=2)
add_table(
    ["Function", "Purpose", "Parameters", "Returns"],
    [
        ["get_fda_recalls(year, limit)",
         "Query openFDA Device Recall API",
         "year (int), limit (int: 1-1000)",
         "DataFrame: recall_number, date, firm, root_cause, product_code, status"],
        ["search_fda_knowledge(query, path)",
         "Search local knowledge base",
         "query (str), path (str)",
         "String of matching lines"],
    ],
)
spacer()

heading("4.4 Technical Details", level=2)
para("API:", bold=True)
bullet("Endpoint: https://api.fda.gov/device/recall.json")
bullet("Auth: Optional API_KEY env variable (higher rate limits)")

para("LLM:", bold=True)
bullet("Ollama local at http://localhost:11434")
bullet("Model: smollm2:1.7b")

para("Packages:", bold=True)
add_table(
    ["Package", "Purpose"],
    [
        ["requests", "HTTP requests to FDA API and Ollama"],
        ["pandas", "Data manipulation and DataFrames"],
        ["tabulate", "df.to_markdown() for table output"],
    ],
)
spacer()

heading("4.5 Usage Instructions", level=2)
para("1. Install dependencies:", bold=True)
doc.add_paragraph("pip install requests pandas tabulate", style="No Spacing")
spacer()
para("2. Start Ollama:", bold=True)
doc.add_paragraph("ollama serve", style="No Spacing")
doc.add_paragraph("ollama pull smollm2:1.7b", style="No Spacing")
spacer()
para("3. Run the main system:", bold=True)
doc.add_paragraph("cd 08_function_calling", style="No Spacing")
doc.add_paragraph("python3 homework2_fda_agent.py", style="No Spacing")

# ── Save ──────────────────────────────────────────────────

out_path = os.path.join(os.path.dirname(__file__), "Homework2_Anjali_Katta.docx")
doc.save(out_path)
print(f"✅ Saved: {out_path}")
