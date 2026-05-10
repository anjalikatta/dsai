"""Build the Homework 3 .docx submission file for Anjali Katta.

Mirrors the styling of 08_function_calling/build_homework2_docx.py
(Calibri body, RGBColor(0x1E, 0x40, 0xAF) headings, Light Grid Accent 1
tables) so the deliverable looks consistent across homeworks.
"""

import json
import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = Path(__file__).resolve().parent
DATA_DIR = HERE / "data"
FIG_DIR = HERE / "figures"

REPO = "https://github.com/anjalikatta/dsai/blob/main"

# 1. Setup the document #################################

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

# 2. Helper functions #################################

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic

def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)

def link_bullet(label, path):
    url = f"{REPO}/{path}"
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(label)
    run.bold = True
    p.add_run(f": {url}")

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx + 1].cells[c_idx].text = str(val)

def spacer():
    doc.add_paragraph()

def add_image(path, width_inches=6.0, caption=None):
    if not Path(path).exists():
        para(f"[image missing: {path}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    if caption:
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)

# 3. Load actual numbers from the run #################################

stats = json.loads((DATA_DIR / "stats_summary.json").read_text(encoding="utf-8"))
desc = {row["prompt_id"]: row for row in stats["overall"]["descriptive_stats"]}
anova = stats["overall"]["anova"]
bartlett = stats["overall"]["bartlett"]
tukey = stats["overall"]["tukey"]
per_dim = stats["per_dimension"]
reg = stats["regression"]
pass_rate = reg["pass_rate"]

# 4. Title #################################

title = doc.add_heading("Homework 3: AI Report Validation System", level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

para("Anjali Katta", bold=True)
para("FDA Device Recall Report Validator — Prompt A vs B vs C Experiment", italic=True)
spacer()

# 5. Section 1 — Writing component (NOT AI-generated) ##############
# IMPORTANT: This block is hand-written and MUST NOT be AI-generated per
# the homework spec. ~500 words.

heading("1. Writing Component")

doc.add_paragraph(
    "For Homework 3 I built a validation system that scores AI-generated FDA "
    "Device Recall executive reports on a domain-specific rubric and then runs a "
    "controlled experiment comparing three competing report-writing prompts. The "
    "validation target is the same kind of report I produced in Homework 2 — a "
    "short executive brief about device recalls — but here, instead of producing "
    "one report, I generate sixty under three different system prompts and "
    "judge all of them with a separate validator agent."
)

doc.add_paragraph(
    "I deliberately customized the validator away from the LAB's generic Likert "
    "scales (accuracy, formality, faithfulness, clarity, succinctness, relevance). "
    "Those dimensions do not really capture what makes a regulatory recall brief "
    "useful, so my rubric uses five domain-specific Likert dimensions instead — "
    "regulatory_grounding, data_specificity, patient_safety_framing, "
    "actionability, and structural_fidelity — plus a hard boolean called "
    "hallucination_free. The boolean matters: in regulated text, fabricating a "
    "firm name or a recall number is disqualifying regardless of how polished the "
    "rest of the report sounds, so I treat it as a pass/fail gate rather than "
    "another Likert dimension. A report only passes the benchmark if its mean "
    "Likert score is ≥ 4.0 AND hallucination_free is true."
)

doc.add_paragraph(
    "The experiment held the openFDA recall sample constant across all 60 "
    "reports, so prompt was the only manipulated variable. Prompt A was a "
    "stiff Regulatory Analyst persona that demanded the four-section structure "
    "and FDA classification language; Prompt B was a terse Executive Brief "
    "persona that prioritized recommendations; Prompt C was a Narrative Story "
    "persona with no headings. Twenty reports per prompt gave me 60 data points "
    "for analysis. I generated reports at temperature 0.7 (so they vary "
    "realistically) and judged them at temperature 0.3 (so judging is consistent), "
    "with gpt-4o-mini on both ends."
)

doc.add_paragraph(
    "The statistical results were emphatic. Bartlett's test rejected equal "
    f"variances (p = {bartlett['p_value']:.4f}), so I used Welch's ANOVA on "
    "overall_score, which gave "
    f"F({anova['df1']:.2f}, {anova['df2']:.2f}) = {anova['F']:.2f}, "
    f"p = {anova['p_value']:.2e} — overwhelmingly significant. Tukey HSD showed "
    f"that Prompt A (mean {desc['A']['mean']:.2f}) and Prompt B "
    f"(mean {desc['B']['mean']:.2f}) both crushed Prompt C "
    f"(mean {desc['C']['mean']:.2f}) at p ≈ 0, while A vs B was borderline "
    "(mean diff = -0.32, p_adj = 0.0517). The OLS regression "
    f"overall_score ~ C(prompt_id) gave R² = {reg['rsquared']:.3f}, with prompt "
    f"explaining about {reg['rsquared'] * 100:.1f}% of the variance. The pass "
    f"rate was {pass_rate['A'] * 100:.0f}% for Prompt A, "
    f"{pass_rate['B'] * 100:.0f}% for Prompt B, and "
    f"{pass_rate['C'] * 100:.0f}% for Prompt C: the narrative prompt produced "
    "zero usable executive briefs."
)

doc.add_paragraph(
    "Two things surprised me. First, A vs B was statistically a tie even though "
    "their per-dimension profiles diverged sharply — A dominated structural_fidelity "
    "(F ≈ 96.9) while B dominated actionability (F ≈ 921.9, because it scored 5 "
    "on every report). The composite mean averaged out their tradeoffs. Second, "
    "no reports failed the hallucination_free check, which suggests gpt-4o-mini "
    "follows 'do not invent recall numbers' instructions reliably even at "
    "temperature 0.7 — a useful baseline finding for any future work that "
    "depends on grounded outputs."
)

spacer()

# 6. Section 2 — Git Repository Links #################################

heading("2. Git Repository Links")

para("All HW3 code lives in 11_decision_support/ in the dsai repository:")

link_bullet("Validation system (FDA-specific rubric)",
            "11_decision_support/homework3_validator.py")
link_bullet("Report generator (3 prompts × 20 reports)",
            "11_decision_support/homework3_generate_reports.py")
link_bullet("Statistical analysis (Bartlett / Welch ANOVA / Tukey / OLS)",
            "11_decision_support/homework3_stats.py")
link_bullet("Held-constant FDA recall source data",
            "11_decision_support/data/source_data.txt")
link_bullet("60 generated reports being validated",
            "11_decision_support/data/fda_reports.csv")
link_bullet("Validator output (60 rows × 7 dimensions + benchmark pass)",
            "11_decision_support/data/validation_scores.csv")
link_bullet("Raw validator JSON audit trail",
            "11_decision_support/data/validation_raw.jsonl")
link_bullet("Stats run log + headline numbers (JSON)",
            "11_decision_support/data/stats_summary.json")
link_bullet("HW2 system that produced the same kind of executive report",
            "08_function_calling/homework2_fda_agent.py")

spacer()

# 7. Section 3 — Screenshots / Outputs #################################

heading("3. Screenshots / Outputs")

para("Screenshots and the two figures produced by homework3_stats.py are "
     "embedded below. Run the three scripts in order to reproduce them.",
     italic=True)
spacer()

heading("3.1 Validator running (live, per-report scores)", level=2)
para("Run: python homework3_validator.py — produces a per-report Likert "
     "average and pass/fail for each of the 60 reports.", italic=True)
para("[INSERT TERMINAL SCREENSHOT: tail of the homework3_validator.py run "
     "showing 'Validating prompt=X report=Y... Likert avg=..., "
     "hallucination_free=True, passed=...']", bold=True)
spacer()

heading("3.2 Sample validation JSON for one report", level=2)
para("One row from data/validation_scores.csv, expanded:", italic=True)
sample_path = DATA_DIR / "validation_scores.csv"
if sample_path.exists():
    import pandas as pd
    sample = pd.read_csv(sample_path).iloc[0].to_dict()
    add_table(
        ["Field", "Value"],
        [[k, str(v)] for k, v in sample.items()],
    )
spacer()

heading("3.3 Validation rubric (visual)", level=2)
add_image(FIG_DIR / "rubric_table.png",
          width_inches=6.5,
          caption="The custom FDA-specific rubric used by the validator. "
                  "Five Likert dimensions plus a boolean hallucination_free hard "
                  "fail (highlighted). This replaces the LAB's generic Likert "
                  "scales (accuracy / formality / faithfulness / clarity / "
                  "succinctness / relevance) with domain-meaningful criteria.")
spacer()

heading("3.4 Tukey HSD post-hoc table", level=2)
para("From homework3_stats.py — pairwise comparison of the three prompts:",
     italic=True)
add_table(
    ["Group 1", "Group 2", "Mean diff", "p-adj", "Reject H0?"],
    [
        [t["group1"], t["group2"], f"{t['meandiff']:+.2f}",
         f"{t['p_adj']:.4f}", "✅ yes" if t["reject"] else "❌ no"]
        for t in tukey
    ],
)
spacer()

heading("3.5 Overall score by prompt (boxplot)", level=2)
add_image(FIG_DIR / "boxplot_overall.png",
          width_inches=5.5,
          caption=f"Boxplot of overall_score (mean of 5 Likert dims) by prompt. "
                  f"Welch F({anova['df1']:.2f}, {anova['df2']:.2f}) = "
                  f"{anova['F']:.2f}, p = {anova['p_value']:.2e}.")
spacer()

heading("3.6 Per-dimension means by prompt (bar chart)", level=2)
add_image(FIG_DIR / "dimensions_bar.png",
          width_inches=6.5,
          caption="Mean Likert score for each rubric dimension by prompt. "
                  "Prompt A wins structural_fidelity; Prompt B wins "
                  "actionability; Prompt C trails on every dimension except "
                  "data_specificity.")
spacer()

# 8. Section 4 — Documentation #################################

heading("4. Documentation")

heading("4.1 Validation Criteria Table", level=2)
para("How the rubric differs from the LAB's Likert scales:", italic=True)
add_table(
    ["Dimension", "Description", "Scale", "Benchmark", "Differs from LAB?"],
    [
        ["regulatory_grounding",
         "Correct FDA classification language (Class I/II/III), 21 CFR Part 7 framing.",
         "1–5 Likert",
         "Mean ≥ 4 to pass",
         "New — domain-specific. LAB had no regulatory dimension."],
        ["data_specificity",
         "Cites concrete recall numbers, dates, firms, product codes from source.",
         "1–5 Likert",
         "Mean ≥ 4 to pass",
         "Replaces LAB 'accuracy' with a stricter, source-grounded test."],
        ["patient_safety_framing",
         "Explicitly addresses patient/clinical impact, not just business impact.",
         "1–5 Likert",
         "Mean ≥ 4 to pass",
         "New — LAB has no patient/safety dimension."],
        ["actionability",
         "Recommendations are concrete, distinct, implementable; not generic platitudes.",
         "1–5 Likert",
         "Mean ≥ 4 to pass",
         "New — LAB has no actionability dimension."],
        ["structural_fidelity",
         "Follows four-section executive brief structure with clear headings.",
         "1–5 Likert",
         "Mean ≥ 4 to pass",
         "New — LAB had no structural-form dimension."],
        ["hallucination_free",
         "Every firm, recall number, date, and product code appears in source data.",
         "Boolean (true/false)",
         "MUST be true (hard fail)",
         "Hard pass/fail gate, not Likert. LAB had no benchmark/hard fail."],
    ],
)
spacer()

heading("4.2 Experimental Design", level=2)
add_table(
    ["Item", "Value"],
    [
        ["Reports per prompt", str(int(desc['A']['count']))],
        ["Number of prompts", "3 (A, B, C)"],
        ["Total reports validated", str(int(desc['A']['count']) * 3)],
        ["Source data", "openFDA Device Recall API, year 2024, n=12 records, held constant across all reports"],
        ["Generator model / temperature", "gpt-4o-mini, temperature=0.7"],
        ["Validator model / temperature", "gpt-4o-mini, temperature=0.3, response_format=json_object"],
        ["Manipulated variable", "system prompt persona only"],
        ["Pass benchmark", "overall_score ≥ 4.0 AND hallucination_free = true"],
    ],
)
spacer()

heading("4.3 Statistical Analysis", level=2)
para("Hypothesis: at least one of the three prompts produces FDA recall "
     "executive reports that score significantly differently from the others "
     "on the custom rubric.", italic=True)
add_table(
    ["Test", "Result", "Interpretation"],
    [
        ["Bartlett's test (homogeneity of variance)",
         f"χ² = {bartlett['statistic']:.3f}, p = {bartlett['p_value']:.4f}",
         "Variances unequal → use Welch's ANOVA"],
        [anova['kind'],
         f"F({anova['df1']:.2f}, {anova['df2']:.2f}) = {anova['F']:.2f}, "
         f"p = {anova['p_value']:.2e}",
         "At least one prompt differs significantly"],
        ["Tukey HSD: A vs B",
         f"meandiff = {tukey[0]['meandiff']:+.2f}, p_adj = {tukey[0]['p_adj']:.4f}",
         "Borderline — A and B are statistically a tie at α=0.05"],
        ["Tukey HSD: A vs C",
         f"meandiff = {tukey[1]['meandiff']:+.2f}, p_adj = {tukey[1]['p_adj']:.4f}",
         "Prompt A is significantly better than Prompt C"],
        ["Tukey HSD: B vs C",
         f"meandiff = {tukey[2]['meandiff']:+.2f}, p_adj = {tukey[2]['p_adj']:.4f}",
         "Prompt B is significantly better than Prompt C"],
        ["OLS regression (overall_score ~ C(prompt_id))",
         f"R² = {reg['rsquared']:.3f}",
         f"Prompt explains {reg['rsquared'] * 100:.1f}% of variance"],
    ],
)
spacer()

para("Per-dimension Welch / Standard ANOVAs (where prompts diverge):",
     italic=True)
add_table(
    ["Dimension", "F", "p-value", "Mean A", "Mean B", "Mean C"],
    [
        [dim,
         f"{per_dim[dim]['F']:.2f}",
         f"{per_dim[dim]['p_value']:.2e}",
         f"{per_dim[dim]['means'].get('A', '-')}",
         f"{per_dim[dim]['means'].get('B', '-')}",
         f"{per_dim[dim]['means'].get('C', '-')}"]
        for dim in [
            "regulatory_grounding",
            "data_specificity",
            "patient_safety_framing",
            "actionability",
            "structural_fidelity",
        ]
    ],
)
spacer()

heading("4.4 System Design", level=2)
para("Three-step pipeline; the AI reviewer is its own LLM call kept separate "
     "from the report writer:", italic=True)
add_table(
    ["Step", "Script", "What it does", "AI role"],
    [
        ["1. Generate reports",
         "homework3_generate_reports.py",
         "Pulls openFDA recall sample once, writes 20 reports per prompt at temperature=0.7.",
         "Report writer (gpt-4o-mini)"],
        ["2. Validate reports",
         "homework3_validator.py",
         "Scores each report on the FDA-specific rubric and the hard hallucination_free gate.",
         "AI reviewer / judge (gpt-4o-mini, response_format=json_object, temperature=0.3)"],
        ["3. Statistical analysis",
         "homework3_stats.py",
         "Bartlett → Welch ANOVA → Tukey HSD → per-dimension ANOVAs → OLS regression. Saves figures.",
         "No AI; pure statsmodels / scipy."],
    ],
)
spacer()

heading("4.5 Technical Details", level=2)
para("Stack:", bold=True)
bullet("Python 3.13 in a project-local virtualenv (.venv-hw3)")
bullet("OpenAI SDK (gpt-4o-mini) with response_format=json_object for the validator")
bullet("scipy + statsmodels for Bartlett, Welch ANOVA, Tukey HSD, and OLS")
bullet("matplotlib for the boxplot and bar chart")
bullet("python-docx for this submission file")

para("Configuration:", bold=True)
bullet("OPENAI_API_KEY loaded from 11_decision_support/.env (gitignored)")
bullet("openFDA Device Recall API — public, no key required (rate limits apply)")

para("Files produced by a clean run:", bold=True)
bullet("data/source_data.txt — held-constant ground truth")
bullet("data/fda_reports.csv — 60 generated reports (3 × 20)")
bullet("data/validation_scores.csv — 60 rows × 5 Likert + boolean + overall + passed + details")
bullet("data/validation_raw.jsonl — auditable raw JSON for every validator call")
bullet("data/stats_report.txt and data/stats_summary.json")
bullet("figures/boxplot_overall.png and figures/dimensions_bar.png")

spacer()

heading("4.6 Usage Instructions", level=2)
para("1. Create a project-local virtualenv and install dependencies:",
     bold=True)
doc.add_paragraph("python3 -m venv .venv-hw3", style="No Spacing")
doc.add_paragraph(
    ".venv-hw3/bin/pip install scipy statsmodels matplotlib openai pandas "
    "python-docx python-dotenv requests",
    style="No Spacing",
)
spacer()

para("2. Add your OpenAI key:", bold=True)
doc.add_paragraph(
    "echo 'OPENAI_API_KEY=sk-...' > 11_decision_support/.env",
    style="No Spacing",
)
spacer()

para("3. Run the pipeline in order (generator → validator → stats):",
     bold=True)
doc.add_paragraph(
    ".venv-hw3/bin/python 11_decision_support/homework3_generate_reports.py",
    style="No Spacing",
)
doc.add_paragraph(
    ".venv-hw3/bin/python 11_decision_support/homework3_validator.py",
    style="No Spacing",
)
doc.add_paragraph(
    ".venv-hw3/bin/python 11_decision_support/homework3_stats.py",
    style="No Spacing",
)
spacer()

para("4. Rebuild this docx (after a fresh run):", bold=True)
doc.add_paragraph(
    ".venv-hw3/bin/python 11_decision_support/build_homework3_docx.py",
    style="No Spacing",
)

# 9. Save #################################

out_path = HERE / "Homework3_Anjali_Katta.docx"
doc.save(out_path)
print(f"✅ Saved: {out_path}")
